import os
import json
import base64
import re
import time
import threading
from concurrent.futures import ThreadPoolExecutor, as_completed
from typing import List, Dict
from google import genai
from google.genai import types
import docx
import shutil
import glob
from pypdf import PdfReader, PdfWriter # Re-adding import
import random # Add random import

# Configuration
def load_api_key():
    key = os.getenv("GEMINI_API_KEY")
    if key:
        return key
    env_paths = [".env", "next-app/.env", "../../.env", "../../../.env"]
    for path in env_paths:
        abs_path = os.path.join(os.path.dirname(__file__), path)
        if os.path.exists(abs_path):
            with open(abs_path, "r") as f:
                for line in f:
                    if line.startswith("GEMINI_API_KEY="):
                        return line.split("=")[1].strip()
        if os.path.exists(path):
            with open(path, "r") as f:
                for line in f:
                    if line.startswith("GEMINI_API_KEY="):
                        return line.split("=")[1].strip()
    return None

GEMINI_API_KEY = load_api_key()
if not GEMINI_API_KEY:
    raise ValueError("GEMINI_API_KEY not found in environment variables or .env")

# Initialize the new Google GenAI Client
client = genai.Client(api_key=GEMINI_API_KEY)
MODEL_NAME = 'gemini-3-pro-preview'

# Thread lock for safe JSON writing
json_lock = threading.Lock()
MAX_WORKERS = 2

SYSTEM_PROMPT = r"""你是一个高水平数学专家和JSON数据提取器。你的任务是从高考数学试卷中完整提取所有选择题和填空题。

输出必须是一个JSON列表，每个题目必须包含：
1. "question_number": 字符串（如 "1", "13"）
2. "type": "single_choice", "multi_choice" (仅新高考), 或 "fill_in"
3. "type_rank": 数字（该题型中的序号）
4. "category": 选择：logic, complex, function, derivative, trigo_func, trigo_sol, sequence, vector, inequality, line_circle, conic, solid_geometry, probability
5. "content": 题干 (LaTeX数学公式用$ $包裹)
6. "has_figure": 布尔值。请自行判断题目（题干或选项）中是否包含图形。如果有，设为 true。无需保存图形本身。
7. "options": 选择题用[{"label":"A","text":"..."}], 填空为null
8. "score_rule": 分值说明

重要约束：
- **完整性**：必须根据文本内容，完整提取所有的小题（单选题、多选题、填空题）。
- **表格处理**：如果题目中包含表格，务必将其转换为标准的 LaTeX tabular 环境并包含在 content 中。
- **严禁**：严谨提取“解答题”。
- **跳过提取**：“程序框图”或“算法”类的选择题要跳过提取，下一道题的type_rank和question_number都相应+1。
- **数学格式**：使用标准的 LaTeX 语法。所有数学符号、公式、变量必须包裹在单个 $ 符号内（例如：$f(x) = \sin x$）。确保符号内部不需要额外的转义，遵循标准的 LaTeX 编写习惯即可。"""

# Global extraction configuration
MAX_WORKERS = 3

def read_word_file(file_path):
    """Read .docx or .doc (via textutil conversion) and return text with placeholders for images/tables."""
    if file_path.endswith('.doc'):
        # Convert .doc to .docx using textutil on macOS
        docx_path = file_path + "x"
        if not os.path.exists(docx_path):
            print(f"Converting {file_path} to docx...")
            os.system(f'textutil -convert docx "{file_path}"')
        file_path = docx_path
    
    try:
        doc = docx.Document(file_path)
        full_content = []
        
        # 1. 第一步：读取所有段落并识别图片占位符
        for para in doc.paragraphs:
            # 停止在解答题部分
            if "解答题" in para.text:
                full_content.append(para.text)
                break
                
            # 检查段落 XML 中是否包含图形引用
            if 'w:drawing' in para._p.xml or 'w:pict' in para._p.xml:
                full_content.append(para.text + " [IMAGE_PLACEHOLDER]")
            else:
                full_content.append(para.text)
        
        # 2. 第二步：读取所有表格并转换为纯文本占位符
        for table in doc.tables:
            table_lines = ["\n[TABLE_START]"]
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells]
                table_lines.append(" | ".join(row_cells))
            table_lines.append("[TABLE_END]\n")
            full_content.append("\n".join(table_lines))
            
        return '\n'.join(full_content)
    except Exception as e:
        print(f"Error reading Word file {file_path}: {e}")
        return ""

def clean_json_text(text):
    """Attempt to fix common JSON issues from LLM LaTeX output."""
    text = text.strip()
    if text.startswith("```json"):
        text = text[7:-3].strip()
    elif text.startswith("```"):
        text = text[3:-3].strip()
    return text

def clean_question_content(questions):
    """Minimal cleanup: only handle newlines and basic space trimming."""
    if not questions: return questions
    
    for q in questions:
        if "content" in q and isinstance(q["content"], str):
            # Just merge single newlines and trim, trust the LLM for Math
            c = q["content"].replace('\r', '')
            q["content"] = re.sub(r'(?<!\n)\n(?!\n)', ' ', c).strip()
    return questions

def find_pdf_file(base_dir, task_name):
    """Recursively find the PDF file for a given task name."""
    # Build a search pattern. We assume the filename starts with the task name.
    # Note: task_name usually includes "（空白卷）", and the file should too.
    pattern = f"{base_dir}/**/{task_name}.pdf"
    files = glob.glob(pattern, recursive=True)
    if files:
        return files[0]
    return None

def merge_questions(new_questions, target_file):
    if not new_questions:
        return
    
    with json_lock:
        existing_data = []
        if os.path.exists(target_file):
            with open(target_file, "r", encoding="utf-8") as f:
                try:
                    existing_data = json.load(f)
                except:
                    existing_data = []
        
        # Map to track existing questions by ID for upsert
        existing_map = {q["id"]: i for i, q in enumerate(existing_data)}
        added_count = 0
        updated_count = 0
        
        for q in new_questions:
            if q["id"] in existing_map:
                # Update existing question
                idx = existing_map[q["id"]]
                # Preserving existing answer if logic dictates, but we generally trust the new extract
                existing_ans = existing_data[idx].get("answer")
                existing_data[idx] = q
                if existing_ans:
                    existing_data[idx]["answer"] = existing_ans
                updated_count += 1
            else:
                # Add new question
                existing_data.append(q)
                existing_map[q["id"]] = len(existing_data) - 1
                added_count += 1
                
        with open(target_file, "w", encoding="utf-8") as f:
            json.dump(existing_data, f, ensure_ascii=False, indent=2)
            
        print(f"    -> Merged: {added_count} new, {updated_count} updated.")

def extract_relevant_pages(file_path):
    """
    Identify pages containing single_choice/fill_in questions.
    Stops when '解答题' (Answer Questions) section header is found.
    Returns: list of page indices (0-based)
    """
    try:
        reader = PdfReader(file_path)
        relevant_pages = []
        
        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            # Check for major section headers indicating the end of small questions
            # Matches: "三、解答题", "四、解答题", "Part II", etc.
            # Using a simplified check for the Chinese header common in Gaokao
            if re.search(r'(三|四|五|Part\s+II|Part\s+B)[、\.]?\s*解答题', text):
                # If the header is at the very beginning, don't include this page
                # If it's later, the page might still contain some fill-in questions
                # For safety, let's include this page but assume subsequent ones are irrelevant
                relevant_pages.append(i)
                break
            relevant_pages.append(i)
            
        return relevant_pages
    except Exception as e:
        print(f"Error reading PDF pages for {file_path}: {e}")
        return []

def process_page_batch(reader, page_indices, source_name, target_json):
    """
    Process a specific subset of pages from the PDF.
    Creates a temporary PDF, uploads it, and merges extracted questions.
    """
    temp_path = None
    try:
        writer = PdfWriter()
        for idx in page_indices:
            writer.add_page(reader.pages[idx])
            
        temp_name = f"temp_batch_{int(time.time())}_{threading.get_ident()}.pdf"
        temp_dir = os.path.dirname(target_json) # Use a safe writable dir like target_json dir
        temp_path = os.path.join(temp_dir, temp_name)
        
        with open(temp_path, "wb") as f:
            writer.write(f)
            
        print(f"    Processing batch for {source_name}: Pages {[p+1 for p in page_indices]}...")
        
        # Upload
        uploaded_file = client.files.upload(file=temp_path, config={'mime_type': 'application/pdf'})
        while uploaded_file.state.name == 'PROCESSING':
            time.sleep(1)
            uploaded_file = client.files.get(name=uploaded_file.name)
            
        if uploaded_file.state.name == 'FAILED':
            raise Exception(f"Batch upload failed: {uploaded_file.state.name}")

        # Prompt
        sys_instruction = SYSTEM_PROMPT + "\\n注意：当前处理的是试卷的一部分页面，题目序号可能不是从1开始。请提取所有可见的选择题和填空题。"
        user_prompt = f"请提取附件中的选择题和填空题。如果题目跨页或不完整，请尽可能根据上下文补全，或者提取已有的部分。"
        
        questions = None
        MAX_RETRIES = 6  # Increase to 6 retries
        for attempt in range(MAX_RETRIES):
            try:
                response = client.models.generate_content(
                    model=MODEL_NAME,
                    contents=[uploaded_file, user_prompt],
                    config=types.GenerateContentConfig(
                        system_instruction=sys_instruction,
                        response_mime_type='application/json'
                    )
                )
                
                raw_text = response.text.strip()
                cleaned_text = clean_json_text(raw_text)
                if not cleaned_text:
                    questions = []
                    break

                try:
                    questions = json.loads(cleaned_text)
                except:
                    fixed = re.sub(r'\\\\(?![\\\\/bfnrtu"])', r'\\\\\\\\', raw_text)
                    try:
                        questions = json.loads(clean_json_text(fixed))
                    except:
                        pass
                
                if questions is not None:
                    questions = clean_question_content(questions)
                    break
                time.sleep(2)
            except Exception as e:
                # Retry logic handling 503/429
                error_str = str(e)
                if "503" in error_str or "429" in error_str or "overloaded" in error_str:
                    wait_time = (2 ** attempt) * 5 + random.uniform(0, 5) # Exponential backoff
                    print(f"      [Server Overloaded] Batch attempt {attempt+1}/{MAX_RETRIES} failed. Waiting {wait_time:.1f}s...")
                    time.sleep(wait_time)
                else:
                    print(f"      Batch attempt {attempt+1}/{MAX_RETRIES} error: {e}")
                    time.sleep(5)
        
        # Cleanup file
        try:
            client.files.delete(name=uploaded_file.name)
        except:
            pass
            
        # Merge results immediately
        if questions:
            ordered_questions = []
            for q in questions:
                # Add extra ID robustness
                q_num = q.get('question_number', 'unknown')
                q_id = f"{source_name}-{q_num}".replace(" ", "")
                new_q = {"id": q_id, "source": source_name}
                
                # Merge fields
                for k, v in q.items():
                    if k not in new_q:
                        new_q[k] = v
                ordered_questions.append(new_q)
            
            merge_questions(ordered_questions, target_json)
            print(f"    Batch success: {len(ordered_questions)} questions merged.")
            return len(ordered_questions)
        return 0
        
    except Exception as e:
        print(f"    Error processing batch: {e}")
        return 0
    finally:
        if temp_path and os.path.exists(temp_path):
            try:
                os.remove(temp_path)
            except:
                pass

def process_paper(file_path, source_name, target_json):
    """
    Main/Split processing logic:
    1. Identify relevant pages (stop at '解答题').
    2. Batch pages (e.g., 2 pages per batch).
    3. Process each batch and update JSON incrementally.
    """
    try:
        # 1. Identify Pages
        relevant_indices = extract_relevant_pages(file_path)
        if not relevant_indices:
            print(f"  [SKIP] {source_name} - No relevant pages found/PDF read error.")
            return [] # Return empty list as we handle merging internally
            
        print(f"  {source_name}: Processing {len(relevant_indices)} pages {relevant_indices}")
        
        reader = PdfReader(file_path)
        
        # 2. Batch Processing (2 pages per batch)
        # Why 2? Enough context for split questions, but small enough to isolate errors.
        BATCH_SIZE = 2
        total_extracted = 0
        
        for i in range(0, len(relevant_indices), BATCH_SIZE):
            batch_indices = relevant_indices[i : i + BATCH_SIZE]
            count = process_page_batch(reader, batch_indices, source_name, target_json)
            total_extracted += count
            
        # Return a dummy list to satisfy the calling method signature if needed, 
        # but merging is ALREADY DONE incrementally inside process_page_batch
        return [1] * total_extracted 
            
    except Exception as e:
        print(f"Error processing {source_name}: {e}")
        return None

def process_task_from_path(file_path, target_json):
    filename = os.path.basename(file_path)
    # Extract source name: remove extension and (空白卷)
    source_name = filename.replace(".pdf", "").replace("（空白卷）", "")
    
    print(f"Starting task: {source_name} (File: {filename})")
    try:
        # Pass target_json to process_paper for incremental saving
        qs = process_paper(file_path, source_name, target_json)
        if qs:
            # Merging already done inside process_paper -> process_page_batch
            return f"  [DONE] {source_name} - Extracted approx {len(qs)} questions (batches merged)"
        return f"  [WARNING] {source_name} - No questions found"
    except Exception as e:
        return f"  [ERROR] {source_name} - {str(e)}"

if __name__ == "__main__":
    # Base directory for Math papers (parent of scripts/)
    BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    # Use absolute path for target JSON as requested
    TARGET_JSON = "/Users/yeatom/VSCodeProjects/gaokao/next-app/src/data/math/small_questions.json"
    
    # Ensure directory exists
    os.makedirs(os.path.dirname(TARGET_JSON), exist_ok=True)
    
    print(f"Scanning for PDF files in {BASE_DIR}...")
    
    # 1. Find all blank papers recursively
    # We are looking for files ending with '（空白卷）.pdf'
    # Pattern: **/*（空白卷）.pdf
    search_pattern = os.path.join(BASE_DIR, "**", "*（空白卷）.pdf")
    all_files = glob.glob(search_pattern, recursive=True)
    
    # Filter out temp files if any
    pdf_files = [f for f in all_files if "temp_batch" not in f]
    
    print(f"Found {len(pdf_files)} PDF files to process.")
    
    # 2. Process in parallel
    with ThreadPoolExecutor(max_workers=MAX_WORKERS) as executor:
        # Submit tasks
        future_to_file = {
            executor.submit(process_task_from_path, pdf_path, TARGET_JSON): pdf_path 
            for pdf_path in pdf_files
        }
        
        for future in as_completed(future_to_file):
            file_path = future_to_file[future]
            try:
                result = future.result()
                print(result)
            except Exception as exc:
                print(f"{os.path.basename(file_path)} generated an exception: {exc}")
